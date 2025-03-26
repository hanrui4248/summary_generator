import streamlit as st
import os
import base64
import datetime as dt
from arxiv_pdf import fetch_papers
from paper_affiliation_classifier import PaperAffiliationClassifier
from paper_assistant import PaperAssistant
import affiliation_analyzer
from abstract_matcher import AbstractMatcher
import re
import pandas as pd
from orgs import orgs
from docx import Document
from docx.shared import Inches
import io
import shutil


def get_download_link(content, filename, text):
    """生成下载链接"""
    b64 = base64.b64encode(content.encode()).decode()
    href = f'<a href="data:text/markdown;base64,{b64}" download="{filename}">{text}</a>'
    return href

def get_binary_file_downloader_html(bin_file, file_label='文件', file_name='文件.docx'):
    """生成二进制文件下载链接"""
    b64 = base64.b64encode(bin_file).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">{file_label}</a>'
    return href

def display_image(image_path, base_dir):
    """显示图片，支持多种路径尝试"""
    # 如果是URL，直接显示
    if image_path.startswith(('http://', 'https://')):
        try:
            st.image(image_path)
            return True
        except Exception as e:
            st.warning(f"无法加载图片URL: {image_path}，错误: {str(e)}")
            return False
    
    # 如果是本地路径，尝试多种路径组合
    possible_paths = [
        image_path,  # 原始路径
        os.path.join(base_dir, image_path),  # 相对于基础目录
        os.path.abspath(image_path),  # 绝对路径
        os.path.join(os.getcwd(), image_path)  # 相对于当前工作目录
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            try:
                st.image(path)
                return True
            except Exception as e:
                st.warning(f"图片存在但无法显示: {path}，错误: {str(e)}")
                break
    
    # st.warning(f"无法找到图片: {image_path}")
    return False

def process_markdown_line(line, base_dir):
    """处理普通的Markdown行，显示文本和图片"""
    # 查找行中的图片标记
    img_pattern = r'!\[(.*?)\]\((.*?)\)'
    img_matches = re.findall(img_pattern, line)
    
    if img_matches:
        # 有图片，分割文本和图片
        parts = re.split(img_pattern, line)
        for i in range(len(parts)):
            if i % 3 == 0:  # 文本部分
                if parts[i].strip():
                    st.markdown(parts[i])
            elif i % 3 == 2:  # 图片路径
                display_image(parts[i], base_dir)
    else:
        # 没有图片，直接显示文本
        st.markdown(line)

def process_markdown_table(table_lines, base_dir):
    """处理Markdown表格，正确显示表格结构和图片"""
    # 解析表格结构
    rows = []
    for line in table_lines:
        # 跳过分隔行 (| --- | --- |)
        if re.match(r'\|\s*[-:]+\s*\|', line):
            continue
        
        # 分割单元格
        cells = line.split('|')[1:-1]  # 去掉首尾的 |
        processed_cells = []
        
        for cell in cells:
            # 检查单元格是否包含图片
            img_match = re.search(r'!\[(.*?)\]\((.*?)\)', cell)
            if img_match:
                # 这是一个包含图片的单元格，使用图片路径作为占位符
                img_path = img_match.group(2)
                processed_cells.append(img_path)
            else:
                # 普通文本单元格
                processed_cells.append(cell.strip())
        
        rows.append(processed_cells)
    
    # 创建表格
    if rows:
        # 创建一个空的DataFrame
        df = pd.DataFrame(rows)
        
        # 使用st.columns创建表格布局
        cols = st.columns(len(rows[0]))
        
        # 填充表格内容
        for col_idx, col in enumerate(cols):
            for row_idx in range(len(rows)):
                cell_content = rows[row_idx][col_idx]
                
                # 检查是否是图片路径
                if re.match(r'.*\.(png|jpg|jpeg|gif|bmp|webp)', cell_content, re.IGNORECASE):
                    with col:
                        display_image(cell_content, base_dir)
                else:
                    with col:
                        st.markdown(cell_content)

def display_markdown_with_images(markdown_content, base_dir="."):
    """解析Markdown内容并使用Streamlit组件显示，包括图片和表格"""
    # 分割内容为行
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        # 检查是否是表格开始
        if lines[i].strip().startswith('|') and '|' in lines[i][1:]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # 处理表格
            process_markdown_table(table_lines, base_dir)
        else:
            # 处理普通文本行
            if lines[i].strip():
                process_markdown_line(lines[i], base_dir)
            i += 1

def markdown_to_docx(markdown_content, base_dir="."):
    """将Markdown内容转换为Word文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('每日arXiv论文快报', 0)
    
    # 分割内容为行
    lines = markdown_content.split('\n')
    i = 0
    
    # 当前段落
    current_paragraph = None
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
            i += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
            i += 1
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
            i += 1
        # 处理表格
        elif line.startswith('|') and '|' in line[1:]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # 解析表格
            rows = []
            for table_line in table_lines:
                # 跳过分隔行
                if re.match(r'\|\s*[-:]+\s*\|', table_line):
                    continue
                
                # 分割单元格
                cells = table_line.split('|')[1:-1]  # 去掉首尾的 |
                rows.append([cell.strip() for cell in cells])
            
            if rows:
                # 创建表格
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                
                # 填充表格内容
                for row_idx, row in enumerate(rows):
                    for col_idx, cell_content in enumerate(row):
                        # 检查单元格是否包含图片
                        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', cell_content)
                        if img_match:
                            # 这是一个包含图片的单元格
                            img_path = img_match.group(2)
                            cell = table.cell(row_idx, col_idx)
                            
                            # 尝试添加图片
                            try:
                                possible_paths = [
                                    img_path,
                                    os.path.join(base_dir, img_path),
                                    os.path.abspath(img_path),
                                    os.path.join(os.getcwd(), img_path)
                                ]
                                
                                for path in possible_paths:
                                    if os.path.exists(path):
                                        cell.paragraphs[0].add_run().add_picture(path, width=Inches(2.0))
                                        break
                            except Exception as e:
                                cell.text = f"[图片: {img_path}]"
                        else:
                            # 普通文本单元格
                            table.cell(row_idx, col_idx).text = cell_content
        
        # 处理图片
        elif '![' in line and '](' in line:
            img_pattern = r'!\[(.*?)\]\((.*?)\)'
            text_parts = re.split(img_pattern, line)
            
            # 创建新段落
            current_paragraph = doc.add_paragraph()
            
            for i_part, part in enumerate(text_parts):
                if i_part % 3 == 0 and part.strip():  # 文本部分
                    current_paragraph.add_run(part)
                elif i_part % 3 == 2:  # 图片路径
                    img_path = part
                    try:
                        possible_paths = [
                            img_path,
                            os.path.join(base_dir, img_path),
                            os.path.abspath(img_path),
                            os.path.join(os.getcwd(), img_path)
                        ]
                        
                        for path in possible_paths:
                            if os.path.exists(path):
                                current_paragraph.add_run().add_picture(path, width=Inches(4.0))
                                break
                    except Exception as e:
                        current_paragraph.add_run(f"[图片: {img_path}]")
            
            i += 1
        
        # 处理普通文本
        elif line:
            doc.add_paragraph(line)
            i += 1
        else:
            # 空行
            i += 1
    
    # 保存到内存中
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io.getvalue()

def main():
    # 清空pdf_folder和images文件夹
    pdf_folder = "pdf_folder"
    if os.path.exists(pdf_folder):
        for file in os.listdir(pdf_folder):
            file_path = os.path.join(pdf_folder, file)
            if os.path.isfile(file_path):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
    else:
        os.makedirs(pdf_folder)
    
    # 清空images文件夹
    images_folder = "images"
    if os.path.exists(images_folder):
        for file in os.listdir(images_folder):
            file_path = os.path.join(images_folder, file)
            if os.path.isfile(file_path):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
    else:
        os.makedirs(images_folder)
    
    st.set_page_config(page_title="每日arXiv论文快报", page_icon="📚")
    
    st.title("AI4Paper")
    st.write("点击下方按钮生成今日精选论文快报")
    
    # 添加一些配置选项
    with st.expander("高级配置"):
        days_back = st.slider("查询过去几天的论文", 1, 7, 1)
        
        # 目标机构多选
        default_orgs = orgs
        target_orgs = st.multiselect(
            "目标机构", 
            options=default_orgs,
            default=default_orgs
        )
        
        # 添加关键词查询选项
        use_keyword_filter = st.checkbox("使用自然语言过滤论文", value=False)
        keyword_query = ""
        if use_keyword_filter:
            keyword_query = st.text_input("请输入提示词", placeholder="例如：强化学习、机器人、大模型等")
    
    # 生成按钮
    if st.button("生成每日论文快报", type="primary"):
        with st.spinner("正在获取和处理论文，请稍候..."):
            # 创建进度条
            progress_bar = st.progress(0)
            
            # 第一步：运行pipeline获取论文,并打上分类标签
            st.info("步骤1/4: 从arXiv获取论文...")
            progress_bar.progress(10)
            
            csv_filename = "papers.csv"
            pdf_folder = "pdf_folder"
            
            # 使用默认的查询字符串 "cat:cs.AI"
            query = "cat:cs.AI"
            
            end_date = dt.datetime.today()
            start_date = end_date - dt.timedelta(days=days_back)
            papers_count = fetch_papers(
                pdf_folder, 
                csv_filename=csv_filename,
                query=query,
                author_filter=False,
                start_date=start_date,
                end_date=end_date
            )
            
            if papers_count == 0:
                st.error("没有找到符合条件的论文，请调整查询参数后重试。")
                return
            
            classifier = PaperAffiliationClassifier()
            classifier.process_csv(csv_filename)
            
            progress_bar.progress(30)
            
            # 第二步：获取目标机构的论文索引
            indices_result = []
            if target_orgs:
                st.info("步骤2/4: 模型筛选目标机构论文...")
                analyzer = affiliation_analyzer.AffiliationAnalyzer()
                indices_result = analyzer.process_csv(csv_filename, target_orgs)
                progress_bar.progress(50)
            
            # 第三步：根据关键词过滤论文
            keyword_indices = []
            if use_keyword_filter and keyword_query:
                st.info("步骤3/4: 根据关键词过滤论文...")
                matcher = AbstractMatcher()
                keyword_indices = matcher.process_csv(csv_filename, keyword_query)
                progress_bar.progress(70)
            else:
                st.info("步骤3/4: 跳过关键词过滤...")
                progress_bar.progress(70)
            
            # 合并过滤结果
            final_indices = []
            if target_orgs and use_keyword_filter and keyword_query:
                # 两种过滤器都启用，取交集
                final_indices = list(set(indices_result).intersection(set(keyword_indices)))
                if not final_indices:
                    st.warning("没有同时满足机构和关键词条件的论文，将显示所有满足机构条件的论文")
                    final_indices = indices_result
            elif target_orgs:
                # 只使用机构过滤
                final_indices = indices_result
            elif use_keyword_filter and keyword_query:
                # 只使用关键词过滤
                final_indices = keyword_indices
            else:
                # 都不使用，获取所有论文索引
                df = pd.read_csv(csv_filename)
                final_indices = list(range(len(df)))
            
            if not final_indices:
                st.error("没有找到符合条件的论文，请调整过滤条件后重试。")
                return
            
            # 第四步：下载论文并生成摘要
            st.info("步骤4/4: 生成论文摘要...")
            assistant = PaperAssistant(output_dir=pdf_folder)
            markdown_content = assistant.process_and_download(csv_filename, final_indices)
            
            progress_bar.progress(100)
            
            # 显示结果
            st.success(f"成功生成论文快报，共包含 {len(final_indices)} 篇论文（从 {papers_count} 篇论文中筛选）")
            # 使用新的display_markdown_with_images函数显示markdown内容
            display_markdown_with_images(markdown_content, pdf_folder)
            
            # 生成Word文档
            docx_content = markdown_to_docx(markdown_content, pdf_folder)
            
            # 提供下载链接
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(
                    get_download_link(markdown_content, "AI生成_每日arXiv精选论文.md", "点击下载Markdown文件"),
                    unsafe_allow_html=True
                )
            with col2:
                st.markdown(
                    get_binary_file_downloader_html(docx_content, "点击下载Word文件", "AI生成_每日arXiv精选论文.md.docx"),
                    unsafe_allow_html=True
                )
            

if __name__ == "__main__":
    main() 