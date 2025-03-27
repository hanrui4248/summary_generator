import base64
import os
import re
import io
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches

def get_download_link(content, filename, text):
    """生成下载链接"""
    b64 = base64.b64encode(content.encode()).decode()
    href = f'<a href="data:text/markdown;base64,{b64}" download="{filename}">{text}</a>'
    return href

def get_binary_file_downloader_html(bin_file, file_label='文件', file_name='文件.docx'):
    """生成二进制文件下载链接"""
    b64 = base64.b64encode(bin_file).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">{file_label}</a>'
    return href

def display_image(image_path, base_dir):
    """显示图片，支持多种路径尝试"""
    # 如果是URL，直接显示
    if image_path.startswith(('http://', 'https://')):
        try:
            st.image(image_path)
            return True
        except Exception as e:
            st.warning(f"无法加载图片URL: {image_path}，错误: {str(e)}")
            return False
    
    # 如果是本地路径，尝试多种路径组合
    possible_paths = [
        image_path,  # 原始路径
        os.path.join(base_dir, image_path),  # 相对于基础目录
        os.path.abspath(image_path),  # 绝对路径
        os.path.join(os.getcwd(), image_path)  # 相对于当前工作目录
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            try:
                st.image(path)
                return True
            except Exception as e:
                st.warning(f"图片存在但无法显示: {path}，错误: {str(e)}")
                break
    
    # st.warning(f"无法找到图片: {image_path}")
    return False

def process_markdown_line(line, base_dir):
    """处理普通的Markdown行，显示文本和图片"""
    # 查找行中的图片标记
    img_pattern = r'!\[(.*?)\]\((.*?)\)'
    img_matches = re.findall(img_pattern, line)
    
    if img_matches:
        # 有图片，分割文本和图片
        parts = re.split(img_pattern, line)
        for i in range(len(parts)):
            if i % 3 == 0:  # 文本部分
                if parts[i].strip():
                    st.markdown(parts[i])
            elif i % 3 == 2:  # 图片路径
                display_image(parts[i], base_dir)
    else:
        # 没有图片，直接显示文本
        st.markdown(line)

def process_markdown_table(table_lines, base_dir):
    """处理Markdown表格，正确显示表格结构和图片"""
    # 解析表格结构
    rows = []
    for line in table_lines:
        # 跳过分隔行 (| --- | --- |)
        if re.match(r'\|\s*[-:]+\s*\|', line):
            continue
        
        # 分割单元格
        cells = line.split('|')[1:-1]  # 去掉首尾的 |
        processed_cells = []
        
        for cell in cells:
            # 检查单元格是否包含图片
            img_match = re.search(r'!\[(.*?)\]\((.*?)\)', cell)
            if img_match:
                # 这是一个包含图片的单元格，使用图片路径作为占位符
                img_path = img_match.group(2)
                processed_cells.append(img_path)
            else:
                # 普通文本单元格
                processed_cells.append(cell.strip())
        
        rows.append(processed_cells)
    
    # 创建表格
    if rows:
        # 创建一个空的DataFrame
        df = pd.DataFrame(rows)
        
        # 使用st.columns创建表格布局
        cols = st.columns(len(rows[0]))
        
        # 填充表格内容
        for col_idx, col in enumerate(cols):
            for row_idx in range(len(rows)):
                cell_content = rows[row_idx][col_idx]
                
                # 检查是否是图片路径
                if re.match(r'.*\.(png|jpg|jpeg|gif|bmp|webp)', cell_content, re.IGNORECASE):
                    with col:
                        display_image(cell_content, base_dir)
                else:
                    with col:
                        st.markdown(cell_content)

def display_markdown_with_images(markdown_content, base_dir="."):
    """解析Markdown内容并使用Streamlit组件显示，包括图片和表格"""
    # 分割内容为行
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        # 检查是否是表格开始
        if lines[i].strip().startswith('|') and '|' in lines[i][1:]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # 处理表格
            process_markdown_table(table_lines, base_dir)
        else:
            # 处理普通文本行
            if lines[i].strip():
                process_markdown_line(lines[i], base_dir)
            i += 1

def markdown_to_docx(markdown_content, base_dir="."):
    """将Markdown内容转换为Word文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('每日arXiv论文快报', 0)
    
    # 分割内容为行
    lines = markdown_content.split('\n')
    i = 0
    
    # 当前段落
    current_paragraph = None
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
            i += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
            i += 1
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
            i += 1
        # 处理表格
        elif line.startswith('|') and '|' in line[1:]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # 解析表格
            rows = []
            for table_line in table_lines:
                # 跳过分隔行
                if re.match(r'\|\s*[-:]+\s*\|', table_line):
                    continue
                
                # 分割单元格
                cells = table_line.split('|')[1:-1]  # 去掉首尾的 |
                rows.append([cell.strip() for cell in cells])
            
            if rows:
                # 创建表格
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                
                # 填充表格内容
                for row_idx, row in enumerate(rows):
                    for col_idx, cell_content in enumerate(row):
                        # 检查单元格是否包含图片
                        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', cell_content)
                        if img_match:
                            # 这是一个包含图片的单元格
                            img_path = img_match.group(2)
                            cell = table.cell(row_idx, col_idx)
                            
                            # 尝试添加图片
                            try:
                                possible_paths = [
                                    img_path,
                                    os.path.join(base_dir, img_path),
                                    os.path.abspath(img_path),
                                    os.path.join(os.getcwd(), img_path)
                                ]
                                
                                for path in possible_paths:
                                    if os.path.exists(path):
                                        cell.paragraphs[0].add_run().add_picture(path, width=Inches(2.0))
                                        break
                            except Exception as e:
                                cell.text = f"[图片: {img_path}]"
                        else:
                            # 普通文本单元格
                            table.cell(row_idx, col_idx).text = cell_content
        
        # 处理图片
        elif '![' in line and '](' in line:
            img_pattern = r'!\[(.*?)\]\((.*?)\)'
            text_parts = re.split(img_pattern, line)
            
            # 创建新段落
            current_paragraph = doc.add_paragraph()
            
            for i_part, part in enumerate(text_parts):
                if i_part % 3 == 0 and part.strip():  # 文本部分
                    current_paragraph.add_run(part)
                elif i_part % 3 == 2:  # 图片路径
                    img_path = part
                    try:
                        possible_paths = [
                            img_path,
                            os.path.join(base_dir, img_path),
                            os.path.abspath(img_path),
                            os.path.join(os.getcwd(), img_path)
                        ]
                        
                        for path in possible_paths:
                            if os.path.exists(path):
                                current_paragraph.add_run().add_picture(path, width=Inches(4.0))
                                break
                    except Exception as e:
                        current_paragraph.add_run(f"[图片: {img_path}]")
            
            i += 1
        
        # 处理普通文本
        elif line:
            doc.add_paragraph(line)
            i += 1
        else:
            # 空行
            i += 1
    
    # 保存到内存中
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io.getvalue()

def load_default_markdown():
    """加载默认的markdown文件"""
    markdown_filename = "每日默认精选论文.md"
    if os.path.exists(markdown_filename):
        try:
            with open(markdown_filename, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            st.warning(f"无法加载默认markdown文件: {str(e)}")
            return None
    else:
        print("未找到默认的论文摘要文件，请点击生成按钮创建")
        return None 